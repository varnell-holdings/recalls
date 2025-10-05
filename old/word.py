from docx import Document
from docx.shared import Pt

import datetime as dt
from jinja2 import Environment, FileSystemLoader


pat = ["Mr Alan MATHISON", "Stoita", "999-999", "Colonoscopy"]
today = dt.date.today()
today_str = today.strftime("%A, %d %B %Y")


def is_over_75(date_of_birth):
    """
    Check if a person is 75 years old or older based on their date of birth.

    Args:
        date_of_birth (str): Date of birth in format "dd/mm/yyyy"

    Returns:
        bool: True if aged 75 or over, False otherwise
    """
    # Parse the date of birth
    dob = dt.datetime.strptime(date_of_birth, "%d/%m/%Y")

    # Get today's date
    today = dt.datetime.now()

    # Calculate age
    age = today.year - dob.year

    # Adjust if birthday hasn't occurred this year yet
    if (today.month, today.day) < (dob.month, dob.day):
        age -= 1


def write_csv():
    pass


def make_letter_text(pat, dob):
    full_name = pat[0]
    title = full_name.split()[0]
    first_name = full_name.split()[1]
    last_name = full_name.split()[-1].title()
    full_name = f"{title} {first_name} {last_name}"

    doctor = pat[1]
    procedure = pat[3]
    if procedure == "COL/PE":
        procedure = "Gastroscopy and Colonoscopy"
    if procedure == "Panendoscopy":
        procedure = "Gastroscopy"

    over_75 = is_over_75(dob)

    path_to_template = "."
    loader = FileSystemLoader(path_to_template)
    env = Environment(loader=loader)
    template_name = "letter_template.txt"
    template = env.get_template(template_name)
    page = template.render(
        today_date=today_str,
        full_name=full_name,
        title=title,
        last_name=last_name,
        doctor=doctor,
        procedure=procedure,
        over_75=over_75,
        ocd=ocd,
    )
    return page

    # with open("letter.txt", "wt") as f:
    #     f.write(page)


def letter_compose():
    today = dt.date.today()
    today_str = today.strftime("%A, %-d %B %Y")
    full_name = pat[0]
    title = full_name.split()[0]
    first_name = full_name.split()[1]
    last_name = full_name.split()[-1].title()
    full_name = f"{title} {first_name} {last_name}"

    page = make_letter_text(pat, dob)

    doc = Document("original_recall_letters/bariol.docx")

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Bookman Old Style"
    font.size = Pt(12)

    paragraph = doc.add_paragraph()

    # Split the text and add line breaks
    lines = page.split("\n")

    for i, line in enumerate(lines):
        print(i, line)
        run = paragraph.add_run(line)
        if i < len(lines) - 1:  # Don't add break after last line
            run.add_break()

    doc.save("current.docx")

    write_csv()


if __name__ == "__main__":
    dob = "12/12/2000"
    ocd = True
    letter_compose()
