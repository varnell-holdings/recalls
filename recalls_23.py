"""This program is used to send second and third recalls to patients at the Diagnostic Endoscopy Centre.
The patient details and date of first recall have been recorded in recalls_csv.csv at the time a the first recall.
It also uses the file episodes.csv which is written to when a patient has a procedure.
On choosing 2nd or 3rd recall from the combobox the program first reads through the csv files and if a patient
has had a procedure since the date of their previous recall message then they are marked as attended in recalls_csv.csv
Then the updated recalls_csv file is read and any patients that are 3 months past the date of their last recall message are added to a list for processing.
Then the program opens the first patient in Blue Chip using their mrn.
The human operator looks in their file and decides to either send a recall message ( email or letter if they don't have an email address)
or they are marked as bo recall required. The program then closes that Blue Chip file using pyautogui and opens the Blue Chip file for the next patient on the list.

"""

import csv
import datetime
import os
from tkinter import FALSE, E, Frame, Menu, N, S, StringVar, Tk, W, messagebox, ttk

# import win32com.client as win32  # pip install pywin32
from dateutil.parser import parse
from docx import Document
from docx.shared import Pt
from jinja2 import Environment, FileSystemLoader

# import pyautogui as pya


# Get today's date
today = datetime.date.today()
today_str = today.strftime("%d-%m-%Y")

# Doctor surname -> first name mapping
doctor_first_names = {
    "Bariol": "Carolyn",
    "Feller": "Robert",
    "Stoita": "Alina",
    "Mill": "Justine",
    "Sanagapalli": "Santosh",
    "Williams": "David",
    "Wettstein": "Antony",
    "Vivekanandarajah": "Suhirdan",
    "Ghaly": "Simon",
    "Vickers": "Chris",
}

# File paths for recalls system
RECALLS_BASE_PATH = "D:\\JOHN TILLET\\source\\active\\recalls"
TEMPLATES_PATH = f"{RECALLS_BASE_PATH}\\templates"
LETTERS_PATH = f"{RECALLS_BASE_PATH}\\letters"
HEADERS_PATH = f"{RECALLS_BASE_PATH}\\headers"
LOGO_PATH = f"{RECALLS_BASE_PATH}\\dec_logo.jpg"

# Screen position for closing Blue Chip window (varies by user/screen)
user = os.getenv("USERNAME")
if user == "John":
    CLOSE_POS = (1020, 120)
elif user == "John2":
    CLOSE_POS = (774, 96)
elif user == "Typing2":
    CLOSE_POS = (780, 96)

recalls_path = "recalls_csv.csv"
episodes_path = "episodes.csv"
patients_to_recall = []
current_record = {}
recall = ""  # 'Second' or 'Third'


def load_recalls(path):
    """Load recalls CSV as dict of dicts keyed by MRN."""
    recalls = {}
    with open(path, "r") as f:
        reader = csv.DictReader(f)
        for row in reader:
            mrn = row["mrn"].strip()
            recalls[mrn] = {
                "name": row["name"].strip(),
                "doctor": row["doctor"].strip(),
                "mrn": mrn,
                "dob": row["dob"].strip(),
                "procedure": row["procedure"].strip(),
                "email": row["email"].strip(),
                "first": row["first"].strip(),
                "second": row["second"].strip(),
                "third": row["third"].strip(),
                "attended": row["attended"].strip(),
            }
    return recalls


def load_episode_dates(path):
    """Load episodes CSV and return dict mapping MRN -> list of dates."""
    episodes = {}
    with open(path, "r") as f:
        reader = csv.DictReader(f)
        for row in reader:
            mrn = row["mrn"].strip()
            date_str = row["date"].strip()
            try:
                episode_date = parse(date_str, dayfirst=True).date()
                if mrn not in episodes:
                    episodes[mrn] = []
                episodes[mrn].append(episode_date)
            except Exception:
                continue
    return episodes


def get_latest_recall_date(record):
    """Get the most recent recall date from first/second columns."""
    dates = []
    for field in ["first", "second"]:
        if record[field]:
            try:
                d = parse(record[field], dayfirst=True).date()
                dates.append(d)
            except Exception:
                continue
    if dates:
        return max(dates)
    return None


def update_recalls_with_episodes(recalls_path, episodes_path):
    """
    Check recalls against episodes and mark patients as attended
    if they've had a procedure after their last recall.

    Returns the updated recalls dict and count of records changed.
    """
    recalls = load_recalls(recalls_path)
    episodes = load_episode_dates(episodes_path)

    updated_count = 0

    for mrn, record in recalls.items():
        # Skip if already attended
        if record["attended"] == "yes":
            continue

        # Get the latest recall date for this patient
        latest_recall = get_latest_recall_date(record)
        if latest_recall is None:
            continue

        # Check if patient has any episode after their recall date
        if mrn in episodes:
            for episode_date in episodes[mrn]:
                if episode_date > latest_recall:
                    record["attended"] = "yes"
                    updated_count += 1
                    print(
                        f"  Updated: {record['name']} (MRN: {mrn}) - recall: {latest_recall}, episode: {episode_date}"
                    )
                    break

    return recalls, updated_count


def save_recalls(recalls, path):
    """Save recalls dict back to CSV."""
    fieldnames = [
        "name",
        "doctor",
        "mrn",
        "dob",
        "procedure",
        "email",
        "first",
        "second",
        "third",
        "attended",
    ]

    with open(path, "w", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        for record in recalls.values():
            writer.writerow(record)


def get_patients_for_second_recall(recalls):
    """
    Filter recalls to find patients needing a second recall.
    Returns list of records where:
    - attended = 'no'
    - first recall date is >12 weeks ago
    - second column is empty
    """
    twelve_weeks_ago = today - datetime.timedelta(weeks=12)
    patients_to_recall = []

    for mrn, record in recalls.items():
        # Skip if already attended
        if record["attended"] == "yes":
            continue

        # Skip if second recall already sent
        if record["second"]:
            continue

        # Skip if no first recall date
        if not record["first"]:
            continue

        # Check if first recall is >12 weeks ago
        try:
            first_date = parse(record["first"], dayfirst=True).date()
            if first_date <= twelve_weeks_ago:
                patients_to_recall.append(record)
        except Exception:
            continue

    return patients_to_recall


def get_procedure_display_name(procedure):
    """Convert procedure code to display name for letters/emails."""
    if procedure == "COL/PE":
        return "Gastroscopy and Colonoscopy"
    if procedure == "Panendoscopy":
        return "Gastroscopy"
    return procedure


def parse_patient_name(full_name):
    """Parse patient name into title, first_name, last_name.

    Assumes format: Title FirstName [MiddleInitial] LASTNAME
    where lastname is in UPPERCASE.
    """
    parts = full_name.split()
    title = parts[0]
    first_name = parts[1]

    # Find where uppercase surname starts (skip title and first name)
    # Also skip single letters (middle initials)
    surname_parts = []
    for part in parts[2:]:
        # Skip middle initials (single letters)
        if len(part) == 1:
            continue
        # Check if word is uppercase (ignoring apostrophes)
        clean = part.replace("'", "")
        if clean.isupper():
            surname_parts.append(part)

    # If no uppercase found, fall back to last word
    if not surname_parts:
        surname_parts = [parts[-1]]

    last_name = " ".join(surname_parts).title()

    return title, first_name, last_name


def open_blue_chip(record):
    """Open Blue Chip CMS with the patient's MRN."""
    pass


def process_csv(event):
    global patients_to_recall
    global current_record
    global recall
    recall = recall_type_var.get()
    if recall == "Second":
        print("Loading and checking recalls against episodes...")
        recalls, updated_count = update_recalls_with_episodes(
            recalls_path, episodes_path
        )

        print(
            f"Found {updated_count} patients who have attended since their last recall."
        )

        save_recalls(recalls, recalls_path)
        print("Saved updated recalls.")

        # Get patients needing second recall
        patients_to_recall = get_patients_for_second_recall(recalls)
        count = len(patients_to_recall)
        print(f"Found {count} patients needing second recall.")
        patient_count_var.set(f"{count} patients to process")

        # Pop first patient to process
        if patients_to_recall:
            current_record = patients_to_recall.pop()
            email = current_record["email"].strip()
            method = "email" if email and "@" in email else "letter"
            current_patient_var.set(f"{current_record['name']} ({method})")
            print(
                f"Processing: {current_record['name']} (MRN: {current_record['mrn']})"
            )
            open_blue_chip(current_record)

    else:
        messagebox.showinfo(message="Can only do second recalls currently")


def write_csv(attended):
    """Update recalls_csv.csv with the action taken on current patient.

    attended='no' - recall sent, write today's date to second/third column
    attended='yes' - no recall needed, mark as attended
    """
    recalls = load_recalls(recalls_path)
    mrn = current_record["mrn"]

    if attended == "no":
        today_str = today.strftime("%d-%m-%Y")
        if recall == "Second":
            recalls[mrn]["second"] = today_str
        else:
            recalls[mrn]["third"] = today_str
    else:
        recalls[mrn]["attended"] = "yes"

    save_recalls(recalls, recalls_path)
    print(f"Updated CSV for {current_record['name']} (MRN: {mrn})")


def make_html_body(our_content_id):
    """Uses jinja2 to construct the html body of the email. Saves in body_.html"""
    full_name = current_record["name"]
    title, first_name, last_name = parse_patient_name(full_name)

    doctor = current_record["doctor"]
    doc_first_name = doctor_first_names[doctor]
    procedure = get_procedure_display_name(current_record["procedure"])

    loader = FileSystemLoader(TEMPLATES_PATH)
    env = Environment(loader=loader)
    if recall == "Second":
        template_name = "email_2_template.html"
    else:
        template_name = "email_3_template.html"
    template = env.get_template(template_name)
    page = template.render(
        today_date=today_str,
        full_name=full_name,
        title=title,
        last_name=last_name,
        doc_first_name=doc_first_name,
        doctor=doctor,
        procedure=procedure,
        our_content_id=our_content_id,
    )

    if recall == "Second":
        body_file = f"{RECALLS_BASE_PATH}\\body_2.html"
    else:
        body_file = f"{RECALLS_BASE_PATH}\\body_3.html"

    with open(body_file, "wt") as f:
        f.write(page)


def email_compose():
    """Create and send recall email via Outlook."""
    outlook = win32.Dispatch("Outlook.Application")
    mail = outlook.CreateItem(0)  # 0 represents an email item
    mail.Subject = "Procedure reminder"

    mail.To = current_record["email"]

    attachment = mail.Attachments.Add(LOGO_PATH)
    CONTENT_ID_PROPERTY = "http://schemas.microsoft.com/mapi/proptag/0x3712001F"
    our_content_id = "my_logo_123"
    attachment.PropertyAccessor.SetProperty(CONTENT_ID_PROPERTY, our_content_id)

    make_html_body(our_content_id)

    if recall == "Second":
        body_file = f"{RECALLS_BASE_PATH}\\body_2.html"
    else:
        body_file = f"{RECALLS_BASE_PATH}\\body_3.html"

    with open(body_file, "rt", encoding="cp1252") as f:
        html_content = f.read()

    mail.HTMLBody = html_content

    mail.Send()

    # write to csv
    write_csv("no")

    next_patient()


def make_letter_text():
    """Generate letter text from template using current_record."""
    full_name = current_record['name']
    title, first_name, last_name = parse_patient_name(full_name)

    doctor = current_record['doctor']
    procedure = get_procedure_display_name(current_record['procedure'])

    loader = FileSystemLoader(TEMPLATES_PATH)
    env = Environment(loader=loader)
    if recall == "Second":
        template_name = "letter_2_template.txt"
    else:
        template_name = "letter_3_template.txt"
    template = env.get_template(template_name)
    page = template.render(
        today_date=today_str,
        full_name=full_name,
        title=title,
        last_name=last_name,
        doctor=doctor,
        procedure=procedure,
    )
    return page


def letter_compose():
    """Create and open a recall letter for current patient."""
    full_name = current_record['name']
    title, first_name, last_name = parse_patient_name(full_name)
    doctor = current_record['doctor']

    page = make_letter_text()

    doc = Document(f"{HEADERS_PATH}\\{doctor}.docx")

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Bookman Old Style"
    font.size = Pt(12)

    paragraph = doc.add_paragraph()

    # Split the text and add line breaks
    lines = page.split("\n")

    for i, line in enumerate(lines):
        run = paragraph.add_run(line)
        if i < len(lines) - 1:  # Don't add break after last line
            run.add_break()

    doc.save(f"{LETTERS_PATH}\\{last_name}.docx")
    write_csv("no")
    os.startfile(f"{LETTERS_PATH}\\{last_name}.docx")
    finish_new_button.grid()


def send_recall():
    """Send recall via email if available, otherwise compose letter."""
    email = current_record["email"].strip()
    if email and "@" in email:
        email_compose()
    else:
        letter_compose()


def no_recall():
    """Mark patient as attended (no recall needed)."""
    write_csv("yes")
    # close_out()
    next_patient()


def close_out():
    """Close the current patient file in Blue Chip using pyautogui."""
    pya.moveTo(CLOSE_POS[0], CLOSE_POS[1])
    pya.click()
    pya.hotkey("alt", "n")


def next_patient():
    """Pop next patient from list and open in Blue Chip."""
    global current_record

    if patients_to_recall:
        current_record = patients_to_recall.pop()
        email = current_record["email"].strip()
        method = "email" if email and "@" in email else "letter"
        current_patient_var.set(f"{current_record['name']} ({method})")
        patient_count_var.set(f"{len(patients_to_recall)} patients to process")
        print(f"Processing: {current_record['name']} (MRN: {current_record['mrn']})")
        open_blue_chip(current_record)
    else:
        current_record = {}
        current_patient_var.set("Finished!")
        patient_count_var.set("0 patients to process")


def finish_recall():
    """Close current patient and open the next one."""
    finish_new_button.grid_remove()
    next_patient()


def open_letters():
    """Open the letters folder in Windows Explorer."""
    os.startfile(LETTERS_PATH)


messagebox.showinfo(message="Make sure Blue Chip is open then press OK.")

root = Tk()

recall_type_var = StringVar()
patient_count_var = StringVar()
current_patient_var = StringVar()

root.geometry("350x450+840+50")
root.title("Second and Third Recalls")
root.option_add("*(tearOff)", FALSE)

menu_bar = Menu(root)
root.config(menu=menu_bar)
extras_menu = Menu(menu_bar)
menu_bar.add_cascade(menu=extras_menu, label="Extras")
extras_menu.add_command(label="Letters folder", command=open_letters)

main_frame = ttk.Frame(root, padding="3 3 12 12")
main_frame.grid(column=0, row=0, sticky=(N, W, E, S))
main_frame.columnconfigure(0, weight=1)
main_frame.rowconfigure(0, weight=1)

top_frame = Frame(main_frame)
top_frame.grid(column=0, row=0, sticky=(N, W, E, S))
top_frame.columnconfigure(0, weight=1)
top_frame.rowconfigure(0, weight=1)

divider_frame = Frame(main_frame, height=2, bg="black")
divider_frame.grid(column=0, row=1, sticky=(N, W, E, S))
divider_frame.columnconfigure(0, weight=1)
divider_frame.rowconfigure(0, weight=1)

bottom_frame = Frame(main_frame)
bottom_frame.grid(column=0, row=2, sticky=(N, W, E, S))
bottom_frame.columnconfigure(0, weight=1)
bottom_frame.rowconfigure(0, weight=1)


recall_type_combo = ttk.Combobox(top_frame, textvariable=recall_type_var, width=30)
recall_type_combo["values"] = ["Second", "Third"]
recall_type_combo.grid(column=0, row=0, sticky=W)
recall_type_combo["state"] = "readonly"
recall_type_combo.bind("<<ComboboxSelected>>", process_csv)

patient_count_label = ttk.Label(top_frame, textvariable=patient_count_var)
patient_count_label.grid(column=0, row=1, sticky=W)

send_recall_button = ttk.Button(bottom_frame, text="Send recall", command=send_recall)
send_recall_button.grid(column=0, row=0, sticky=W)

current_patient_label = ttk.Label(bottom_frame, textvariable=current_patient_var)
current_patient_label.grid(column=1, row=0, sticky=E)

no_recall_button = ttk.Button(bottom_frame, text="No Recall", command=no_recall)
no_recall_button.grid(column=0, row=1, sticky=W)

finish_new_button = ttk.Button(bottom_frame, text="Finish & next", command=finish_recall)
finish_new_button.grid(column=0, row=2, sticky=W)

for child in main_frame.winfo_children():
    child.grid_configure(padx=5, pady=5)

for child in top_frame.winfo_children():
    child.grid_configure(padx=5, pady=10)

for child in bottom_frame.winfo_children():
    child.grid_configure(padx=5, pady=20)

finish_new_button.grid_remove()

recall_type_combo.set("")
patient_count_var.set("Choose type of recall to start!")

root.attributes("-topmost", True)
root.mainloop()
