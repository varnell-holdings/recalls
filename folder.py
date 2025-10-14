# import win32com.client as win32  # pip install pywin32
import argparse
import csv
import datetime
from dateutil.parser import parse
import os
import re
import shutil
import time

from tkinter import ttk, StringVar, Tk, W, E, N, S
from tkinter import FALSE, Menu, Frame, messagebox


from jinja2 import Environment, FileSystemLoader

from docx import Document
from docx.shared import Pt

# import pyautogui as pya
import pyperclip
from pyisemail import is_email

# pya.PAUSE = 0.1

# Get today's date
today = datetime.date.today()

parser = argparse.ArgumentParser(description="This is a test script")
parser.add_argument("-t", "--test", action="store_true", help="Run the test")
args = parser.parse_args()
if args.test:
    print("Test mode activated")
    csv_address = "d:\\john tillet\\source\\active\\recalls\\test_csv.csv"
    csv_address_2 = "D:\\Nobue\\test_recalls_csv.csv"
    disposal_csv_address = "d:\\john tillet\\source\\active\\recalls\\test_disposal.csv"


else:
    print("No test mode activated")
    csv_address = "D:\\JOHN TILLET\\source\\active\\recalls\\recalls_csv.csv"
    csv_address_2 = "D:\\Nobue\\recalls_csv.csv"
    disposal_csv_address = "d:\\john tillet\\source\\active\\recalls\\disposal.csv"


full_path = ""
print_length = 0
pat = []  # eg ['Mr Alan MATHISON', 'Stoita', '9998877', 'Colonoscopy']
email = ""
phone = ""
mrn = ""
dob = ""
recall_type = "none"


ocd_doc_set = {
    "Bariol",
    "Feller",
    "Stoita",
    "Mill",
}
# non_ocd_doc_set = {"Sanagapalli", "Williams",
#                    "Wettstein", "Vivekanandahrajah", "Ghaly"}

DOCTORS = [
    "Bariol",
    "Feller",
    "Ghaly",
    "Mill",
    "Sanagapalli",
    "Stoita",
    "Vivekanandahrajah",
    "Wettstein",
    "Williams",
]


full_doc_dict = {
    "Bariol": "Carolyn",
    "Feller": "Robert",
    "Stoita": "Alina",
    "Mill": "Justine",
    "Sanagapalli": "Santosh",
    "Williams": "David",
    "Wettstein": "Antony",
    "Vivekanandahrajah": "Suhirdan",
    "Ghaly": "Simon",
}

proc_dict = {"Colonoscopy": "c", "COL/PE": "d", "Panendoscopy": "p"}

user = os.getenv("USERNAME")

if user == "John":
    RED_BAR_POS = (280, 790)
    TITLE_POS = (230, 170)
    MRN_POS = (740, 315)
    POST_CODE_POS = (610, 355)
    DOB_POS = (750, 220)
    FUND_NO_POS = (770, 703)
    CLOSE_POS = (1020, 120)
    SMS_POS = (485, 735)
    EMAIL_POS = (380, 510)
elif user == "John2":
    RED_BAR_POS = (160, 630)
    TITLE_POS = (200, 134)
    MRN_POS = (600, 250)
    POST_CODE_POS = (490, 284)
    DOB_POS = (600, 174)
    FUND_NO_POS = (580, 548)
    CLOSE_POS = (774, 96)
    SMS_POS = (360, 630)
    EMAIL_POS = (335, 395)
elif user == "Recept5":
    RED_BAR_POS = (160, 630)
    TITLE_POS = (200, 134)
    MRN_POS = (575, 250)
    POST_CODE_POS = (480, 280)
    DOB_POS = (600, 174)
    FUND_NO_POS = (580, 548)
    CLOSE_POS = (780, 96)
elif user == "Typing2":
    RED_BAR_POS = (160, 630)
    TITLE_POS = (200, 134)
    MRN_POS = (575, 250)
    POST_CODE_POS = (480, 280)
    DOB_POS = (600, 174)
    FUND_NO_POS = (580, 548)
    CLOSE_POS = (780, 96)


def open_bc():
    pass
    # pya.moveTo(100, 450, duration=0.3)
    # pya.click()
    # pya.hotkey("ctrl", "o")


def scraper(email=False):
    """'"""
    pass


#     result = "na"
#     pya.hotkey("ctrl", "c")
#     result = pyperclip.paste()
#     if email:
#         result = re.split(r"[\s,:/;\\]", result)[0]
#         if not is_email(result):
#             result = ""
#     return result


def scrape():
    pass


#     global mrn
#     global email
# global dob

# pya.moveTo(100, 450, duration=0.1)
# pya.click()
# pya.press("up", presses=9)
# pya.press("enter")
# time.sleep(0.5)

# pya.moveTo(EMAIL_POS)
# pya.doubleClick()
# email = scraper()
# print(email)

# pya.moveTo(MRN_POS)
# pya.doubleClick()
# mrn = scraper()
# print(mrn)

# pya.moveTo(DOB_POS)
# pya.doubleClick()
# dob = scraper()
# print(dob)

# if (not mrn.isdigit()) or (not parse(dob)):
#     scrape_info_label.set("Error in data")
#     root.update_idletasks()
#     return

# if is_email(email) and email not in {"", "na"}:
#     scrape_info_label.set("OK")
# else:
#     scrape_info_label.set("Problem with email")
# root.update_idletasks()


def parse_dob():
    try:
        parse(dob, dayfirst=True)
        return True
    except Exception:
        return False


def write_csv():
    day_sent = today.isoformat()
    with open(csv_address, "a") as f:
        writer = csv.writer(f, dialect="excel", lineterminator="\n")
        # name, doctor, phone, mrn, dob, phone, email, first, second, third, gp, attended
        # - first second and third are dates
        entry = [
            pat[0],
            pat[1],
            pat[2],
            mrn,
            dob,
            phone,
            email,
            day_sent,
            "",
            "",
            "no",
            "",
        ]
        writer.writerow(entry)
    shutil.copy(csv_address, csv_address_2)


def make_html_body(our_content_id):
    full_name = pat[0]
    title = full_name.split()[0]
    first_name = full_name.split()[1]
    last_name = full_name.split()[-1].title()
    full_name = f"{title} {first_name} {last_name}"

    doctor = pat[1]

    doc_first_name = full_doc_dict[doctor]

    if doctor in ocd_doc_set:
        ocd = True
    else:
        ocd = False

    procedure = pat[3]
    if procedure == "COL/PE":
        procedure = "Gastroscopy and Colonoscopy"
    if procedure == "Panendoscopy":
        procedure = "Gastroscopy"

    today_str = today.strftime("%d-%m-%Y")

    path_to_template = "D:\\JOHN TILLET\\source\\active\\recalls"
    loader = FileSystemLoader(path_to_template)
    env = Environment(loader=loader)
    template_name = "body_1_template.html"
    template = env.get_template(template_name)
    page = template.render(
        today_date=today_str,
        full_name=full_name,
        title=title,
        last_name=last_name,
        doc_first_name=doc_first_name,
        doctor=doctor,
        procedure=procedure,
        over_75=over_75,
        ocd=ocd,
        our_content_id=our_content_id,
    )

    with open("D:\\JOHN TILLET\\source\\active\\recalls\\body_2.html", "wt") as f:
        f.write(page)


def letter_compose():
    global recall_type
    recall_type = "letter"
    full_name = pat[0]
    title = full_name.split()[0]
    first_name = full_name.split()[1]
    last_name = full_name.split()[-1].title()
    full_name = f"{title} {first_name} {last_name}"
    doctor = pat[1]

    page = make_letter_text(pat, dob)

    doc = Document(f"d:\\john tillet\\source\\active\\recalls\\headers\\{doctor}.docx")

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Bookman Old Style"
    font.size = Pt(12)

    paragraph = doc.add_paragraph()

    # Split the text and add line breaks
    lines = page.split("\n")

    for i, line in enumerate(lines):
        print(i, line)
        run = paragraph.add_run(line)
        if i < len(lines) - 1:  # Don't add break after last line
            run.add_break()

    doc.save(f"d:\\john tillet\\source\\active\\recalls\\letters\\{last_name}.docx")
    write_csv()
    scrape_info_label.set("Letter made")
    root.update_idletasks()


# def recall_type(event):
#     recall_type = rec.get()
#     if recall_type == "GP":
#         button4.config(state="disabled", style="Disabled.TButton")
#         send_text_button.config(state="disabled", style="Disabled.TButton")
#     else:
#         button4.config(state="normal", style="Normal.TButton")
#         send_text_button.config(state="normal", style="Normal.TButton")
#     root.update_idletasks()


def recall_compose():
    outlook = win32.Dispatch("Outlook.Application")
    mail = outlook.CreateItem(0)  # 0 represents an email item
    mail.Subject = "Procedure reminder"

    mail.To = email

    logo_path = r"D:\\JOHN TILLET\\source\\active\\recalls\\dec_logo.jpg"
    attachment = mail.Attachments.Add(logo_path)
    CONTENT_ID_PROPERTY = "http://schemas.microsoft.com/mapi/proptag/0x3712001F"
    our_content_id = "my_logo_123"
    attachment.PropertyAccessor.SetProperty(CONTENT_ID_PROPERTY, our_content_id)

    make_html_body(our_content_id)

    with open(
        "D:\\JOHN TILLET\\source\\active\\recalls\\body_2.html", "rt", encoding="cp1252"
    ) as f:
        html_content = f.read()

    mail.HTMLBody = html_content
    # # Uncomment to actually send the email
    # # mail.Send()

    # # Or display it for review before sending
    mail.Display()
    # write to csv
    day_sent = today.isoformat()
    with open(csv_address, "a") as f:
        writer = csv.writer(f, dialect="excel", lineterminator="\n")
        # name, doctor, phone, mrn, dob, phone, email, first, second, third, gp, attended
        # - first second and third are dates
        entry = [
            pat[0],
            pat[1],
            pat[2],
            mrn,
            dob,
            phone,
            email,
            day_sent,
            "",
            "",
            "no",
            "",
        ]
        writer.writerow(entry)
    shutil.copy(csv_address, csv_address_2)


def no_recall():
    pass
    proc.set("Procedure")
    # close bc


def finish():
    """This will close Blue Chip and write to
    disposal.csv and reset buttons above the line"""
    global recall_type
    day_sent = today.isoformat()
    full_name = ""  # get this from scraper
    doctor = doc.get()
    recall_number = rec.get()
    with open(disposal_csv_address, "a") as f:
        writer = csv.writer(f)
        entry = (day_sent, full_name, doctor, recall_number, recall_type)
        writer.writerow(entry)
    recall_type = "none"
    # doc.set("Doctor")
    # rec.set("Recall Type")
    proc.set("Procedure")


root = Tk()

doc = StringVar()  # doctor for recall
rec = StringVar()  # number of recall ie 2 or 3
proc = StringVar()  # type of procedure
scrape_info_label = StringVar()

root.geometry("450x550+840+50")
root.title("Recalls from Folder")
root.option_add("*(tearOff)", FALSE)

# Create a style
style = ttk.Style()

# Configure styles for different states
style.configure("Normal.TButton", background="lightgray", foreground="blue")
style.configure("Disabled.TButton", background="lightgray", foreground="darkgray")

mainframe = ttk.Frame(root, padding="3 3 12 12")
mainframe.grid(column=0, row=0, sticky=(N, W, E, S))
mainframe.columnconfigure(0, weight=1)
mainframe.rowconfigure(0, weight=1)

topframe = Frame(mainframe)
topframe.grid(column=0, row=0, sticky=(N, W, E, S))
topframe.columnconfigure(0, weight=1)
topframe.rowconfigure(0, weight=1)

midframe = Frame(mainframe, height=2, bg="black")
midframe.grid(column=0, row=2, sticky=(N, W, E, S))
midframe.columnconfigure(0, weight=1)
midframe.rowconfigure(0, weight=1)

bottomframe = Frame(mainframe)
bottomframe.grid(column=0, row=3, sticky=(N, W, E, S))
bottomframe.columnconfigure(0, weight=1)
bottomframe.rowconfigure(0, weight=1)


doc_box = ttk.Combobox(topframe, textvariable=doc)
doc_box["values"] = DOCTORS
doc_box["state"] = "readonly"
doc_box.grid(column=0, row=0, sticky=W)

rec_box = ttk.Combobox(topframe, textvariable=rec)
rec_box["values"] = ["second", "third"]
rec_box["state"] = "readonly"
rec_box.grid(column=0, row=1, sticky=W)

proc_box = ttk.Combobox(topframe, textvariable=proc)
proc_box["values"] = ["Pe", "Col", "Double"]
proc_box["state"] = "readonly"
proc_box.grid(column=1, row=1, sticky=W)

open_button = ttk.Button(bottomframe, text="Open Blue Chip", command=open_bc)
open_button.grid(column=0, row=0, sticky=W)


scrape_button = ttk.Button(bottomframe, text="Get info", command=scrape)
scrape_button.grid(column=0, row=1, sticky=W)

scrape_label = ttk.Label(bottomframe, textvariable=scrape_info_label)
scrape_label.grid(column=1, row=1, sticky=E)


button4 = ttk.Button(bottomframe, text="Send email", command=recall_compose)
button4.grid(column=0, row=2, sticky=W)

letter_button = ttk.Button(
    bottomframe, text="Make Recall letter", command=letter_compose
)
letter_button.grid(column=1, row=2, sticky=E)

no_recall_button = ttk.Button(bottomframe, text="No recall", command=no_recall)
no_recall_button.grid(column=1, row=3, sticky=W)


finish_button = ttk.Button(bottomframe, text="Finish patient", command=finish)
finish_button.grid(column=0, row=3, sticky=W)

for child in topframe.winfo_children():
    child.grid_configure(padx=5, pady=20)

for child in bottomframe.winfo_children():
    child.grid_configure(padx=5, pady=20)


scrape_info_label.set("")

doc.set("Doctor")
rec.set("Recall Number")
proc.set("Procedure")

root.attributes("-topmost", True)
root.mainloop()
