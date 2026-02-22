import argparse
import csv
import datetime
import os
import pickle
import re
import shutil
import subprocess
import sys
import time
from pathlib import Path
from tkinter import FALSE, E, Frame, Menu, N, S, StringVar, Tk, W, ttk
from tkinter.filedialog import askopenfilename

import pyautogui as pya
import pyperclip
import win32com.client as win32  # pip install pywin32
from dateutil.parser import parse
from docx import Document
from docx.shared import Pt
from jinja2 import Environment, FileSystemLoader
from pyisemail import is_email
from striprtf.striprtf import rtf_to_text

pya.PAUSE = 0.1

# --- Dates ---
today = datetime.date.today()
today_str = today.strftime("%d-%m-%Y")

# --- Command line arguments ---
parser = argparse.ArgumentParser(description="Recalls - First")
parser.add_argument("-t", "--test", action="store_true", help="Run the test")
parser.add_argument("-n", "--nopickle", action="store_true", help="No pickle option")
args = parser.parse_args()

# --- Paths ---
BASE_PATH = Path("D:/JOHN TILLET/source/active/recalls")
TEMPLATE_PATH = BASE_PATH / "templates"
LETTERS_PATH = BASE_PATH / "letters"
HEADERS_PATH = BASE_PATH / "headers"
BODY_HTML_PATH = BASE_PATH / "body_1.html"
LOGO_PATH = BASE_PATH / "dec_logo.jpg"
OLD_FILES_PATH = BASE_PATH / "old"
PICKLER_PATH = BASE_PATH / "pickler.py"

if args.test:
    csv_address = Path("d:/john tillet/source/active/recalls/csv/test_csv.csv")
    csv_address_2 = Path("D:/Nobue/test_recalls_csv.csv")
else:
    csv_address = BASE_PATH / "csv" / "recalls_csv.csv"
    csv_address_2 = Path("D:/Nobue/recalls_csv.csv")

if not args.nopickle:
    pickle_address = BASE_PATH / "pickled_list"

# --- Shared state (replaces individual global variables) ---
state = {
    "full_path": "",
    "num_to_do": 0,
    "pat": [],  # eg ['Mr Alan MATHISON', 'Stoita', '0432-876-980', 'Colonoscopy']
    "email": "",
    "phone": "",
    "mrn": "",
    "dob": "",
    "first_run": True,
    "output_list_4": [],
    "recall_type": "",
    "recall_number": "",
}

# --- GUI widgets (populated in main, used by callbacks) ---
widgets = {}

# --- Doctor data ---
ocd_doc_set = {
    "Bariol",
    "Feller",
    "Stoita",
    "Mill",
}

full_doc_dict = {
    "Bariol": "Carolyn",
    "Feller": "Robert",
    "Stoita": "Alina",
    "Mill": "Justine",
    "Sanagapalli": "Santosh",
    "Williams": "David",
    "Wettstein": "Antony",
    "Vivekanandarajah": "Suhirdan",
    "Ghaly": "Simon",
    "Vickers": "Chris",
}

proc_dict = {"Colonoscopy": "c", "COL/PE": "d", "Panendoscopy": "p"}

# --- Screen positions (vary by workstation) ---
user = os.getenv("USERNAME")

if user == "John":
    RED_BAR_POS = (280, 790)
    TITLE_POS = (230, 170)
    MRN_POS = (740, 315)
    POST_CODE_POS = (610, 355)
    DOB_POS = (750, 220)
    FUND_NO_POS = (770, 703)
    CLOSE_POS = (1020, 120)
    SMS_POS = (485, 735)
    EMAIL_POS = (600, 510)
elif user == "John2":
    RED_BAR_POS = (160, 630)
    TITLE_POS = (200, 134)
    MRN_POS = (600, 250)
    POST_CODE_POS = (490, 284)
    DOB_POS = (600, 174)
    FUND_NO_POS = (580, 548)
    CLOSE_POS = (774, 96)
    SMS_POS = (360, 630)
    EMAIL_POS = (450, 395)
elif user == "Typing2":
    RED_BAR_POS = (160, 630)
    TITLE_POS = (200, 134)
    MRN_POS = (575, 250)
    POST_CODE_POS = (480, 280)
    DOB_POS = (600, 174)
    FUND_NO_POS = (580, 548)
    CLOSE_POS = (780, 96)
    SMS_POS = (360, 630)
    EMAIL_POS = (450, 400)


# --- Pure utility helpers ---


def parse_patient_name(pat):
    """Extract title, first name, last name from patient record."""
    full_name = pat[0]
    title = full_name.split()[0]
    first_name = full_name.split()[1]
    last_name = full_name.split()[-1].title()
    full_name_formatted = f"{title} {first_name} {last_name}"
    return title, first_name, last_name, full_name_formatted


def postcode_to_state(postcode):
    """Convert Australian postcode to state abbreviation."""
    post_dic = {"3": "VIC", "4": "QLD", "5": "SA", "6": "WA", "7": "TAS"}
    try:
        if postcode[0] == "0":
            if postcode[:2] in {"08", "09"}:
                return "NT"
            else:
                return ""
        elif postcode[0] in {"0", "1", "8", "9"}:
            return ""
        elif postcode[0] == "2":
            if (2600 <= int(postcode) <= 2618) or postcode[:2] == "29":
                return "ACT"
            else:
                return "NSW"
        else:
            return post_dic[postcode[0]]
    except Exception:
        return ""


def parse_dob():
    try:
        parse(state["dob"], dayfirst=True)
        return True
    except Exception:
        return False


def is_over_75(date_of_birth):
    dob = datetime.datetime.strptime(date_of_birth, "%d/%m/%Y")
    age = today.year - dob.year
    if (today.month, today.day) < (dob.month, dob.day):
        age -= 1
    return age > 75


# --- Data persistence ---


def get_pickled_list():
    with open(pickle_address, "rb") as fh:
        state["output_list_4"] = pickle.load(fh)


def set_pickled_list():
    with open(pickle_address, "wb") as fh:
        pickle.dump(state["output_list_4"], fh)


def write_csv(attended):
    day_sent = today.isoformat()
    name = state["pat"][0]
    doctor = state["pat"][1]
    procedure = state["pat"][3]
    email_to_write = state["email"]  # may be email or pipe-separated address
    if attended == "yes":
        first = ""
    else:
        first = day_sent
    with open(csv_address, "a") as fh:
        writer = csv.writer(fh, dialect="excel", lineterminator="\n")
        # name, doctor, mrn, dob, procedure, email, first, second, third, attended
        entry = [
            name,
            doctor,
            state["mrn"],
            state["dob"],
            procedure,
            email_to_write,
            first,
            "",
            "",
            attended,
        ]
        writer.writerow(entry)
    shutil.copy(csv_address, csv_address_2)


# --- Blue Chip interaction (scraping and navigation) ---


def scraper(email=False):
    result = "na"
    pya.hotkey("ctrl", "c")
    result = pyperclip.paste()
    if email:
        result = re.split(r"[\s,:/;\\*]", result)[0]
        for object in result:
            if is_email(object):
                return object
            return ""
    return result


def address_scrape():
    """Scrape street, suburb, postcode from Blue Chip patient info page."""
    pya.moveTo(100, 450, duration=0.1)
    pya.click()
    pya.hotkey("alt", "b")
    pya.press("tab", presses=3)
    street = scraper()
    street = street.replace(",", "")

    pya.press("tab")
    pya.press("tab")
    suburb = scraper()

    pya.moveTo(POST_CODE_POS, duration=0.1)
    pya.doubleClick()
    postcode = scraper()

    aus_state = postcode_to_state(postcode)

    address1 = f"{street}"
    address2 = f"{suburb} {aus_state} {postcode}"

    return address1, address2


def open_bc_by_name():
    name_as_list = state["pat"][0].split()
    name_for_bc = name_as_list[-1] + "," + name_as_list[1]
    pya.moveTo(100, 450, duration=0.3)
    pya.click()
    pya.hotkey("ctrl", "o")
    pya.typewrite(name_for_bc)
    pya.press("enter")


def open_bc_by_phone():
    phone = state["pat"][2].replace("-", "")
    pya.moveTo(100, 450, duration=0.3)
    pya.click()
    pya.hotkey("ctrl", "o")
    pya.hotkey("alt", "b")
    time.sleep(0.2)

    pya.press("down", presses=4)

    pya.hotkey("shift", "tab")
    pya.hotkey("shift", "tab")
    pya.typewrite(phone)
    pya.press("enter")
    pya.moveTo(100, 450, duration=0.3)
    pya.click()
    pya.press("up", presses=4)
    pya.press("enter")


# --- Content generation (email and letter templates) ---


def make_html_body(our_content_id):
    """Uses jinja2 to construct the html body of the email. Saves in body_1.html"""
    title, first_name, last_name, full_name = parse_patient_name(state["pat"])
    doctor = state["pat"][1]
    doc_first_name = full_doc_dict[doctor]
    ocd = doctor in ocd_doc_set

    procedure = state["pat"][3]
    if procedure == "COL/PE":
        procedure = "Gastroscopy and Colonoscopy"
    if procedure == "Panendoscopy":
        procedure = "Gastroscopy"

    over_75 = is_over_75(state["dob"])

    loader = FileSystemLoader(TEMPLATE_PATH)
    env = Environment(loader=loader)
    template = env.get_template("email_1_template.html")
    page = template.render(
        today_date=today_str,
        full_name=full_name,
        title=title,
        last_name=last_name,
        doc_first_name=doc_first_name,
        doctor=doctor,
        procedure=procedure,
        over_75=over_75,
        ocd=ocd,
        our_content_id=our_content_id,
    )

    BODY_HTML_PATH.write_text(page)


def make_letter_text(pat, dob, address1, address2):
    title, first_name, last_name, full_name = parse_patient_name(pat)
    doctor = pat[1]
    procedure = pat[3]
    if procedure == "COL/PE":
        procedure = "Gastroscopy and Colonoscopy"
    if procedure == "Panendoscopy":
        procedure = "Gastroscopy"

    over_75 = is_over_75(dob)
    ocd = doctor in ocd_doc_set

    loader = FileSystemLoader(TEMPLATE_PATH)
    env = Environment(loader=loader)
    template = env.get_template("letter_1_template.txt")
    page = template.render(
        today_date=today_str,
        full_name=full_name,
        title=title,
        last_name=last_name,
        address1=address1,
        address2=address2,
        doctor=doctor,
        procedure=procedure,
        over_75=over_75,
        ocd=ocd,
    )
    return page


# --- Workflow actions (send text, email, letter) ---


def send_text():
    widgets["patient_label_var"].set("")
    title, first_name, last_name, full_name = parse_patient_name(state["pat"])
    doctor = state["pat"][1]
    message = f"Dear {title} {
        last_name
    } just advising you that an email will be sent to you from Dr {
        doctor
    } with a reminder that you are now due for your procedure. Please review email and contact our office on 83826622 if you have any queries and for all bookings."

    pya.moveTo(100, 450, duration=0.3)
    pya.click()
    pya.press("up", presses=3)
    pya.press("enter")
    pya.hotkey("alt", "n")
    pya.moveTo(SMS_POS[0], SMS_POS[1])
    pya.click()
    pya.typewrite(message)
    time.sleep(2)
    pya.press("tab")
    pya.press("enter")
    time.sleep(3)
    pya.press("enter")
    widgets["scrape_info_label"].set("Text sent - open Outlook")
    widgets["root"].update_idletasks()


def recall_compose():
    outlook = win32.Dispatch("Outlook.Application")
    mail = outlook.CreateItem(0)
    mail.Subject = "Procedure reminder"

    mail.To = state["email"]

    attachment = mail.Attachments.Add(str(LOGO_PATH))
    CONTENT_ID_PROPERTY = "http://schemas.microsoft.com/mapi/proptag/0x3712001F"
    our_content_id = "my_logo_123"
    attachment.PropertyAccessor.SetProperty(CONTENT_ID_PROPERTY, our_content_id)

    make_html_body(our_content_id)

    html_content = BODY_HTML_PATH.read_text(encoding="cp1252")

    mail.HTMLBody = html_content
    mail.Send()
    write_csv(attended="no")

    widgets["scrape_info_label"].set("Email made")
    widgets["root"].update_idletasks()


def letter_compose():
    state["recall_type"] = "letter"
    title, first_name, last_name, full_name = parse_patient_name(state["pat"])
    doctor = state["pat"][1]

    address1, address2 = address_scrape()
    state["email"] = f"{address1} | {address2}"

    page = make_letter_text(state["pat"], state["dob"], address1, address2)

    doc = Document(HEADERS_PATH / f"{doctor}.docx")

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Bookman Old Style"
    font.size = Pt(12)

    paragraph = doc.add_paragraph()

    lines = page.split("\n")

    for i, line in enumerate(lines):
        run = paragraph.add_run(line)
        if i < len(lines) - 1:
            run.add_break()

    folder = LETTERS_PATH / today.isoformat()
    folder.mkdir(parents=True, exist_ok=True)
    letter_path = folder / f"{last_name}.docx"
    doc.save(letter_path)
    write_csv(attended="no")
    widgets["scrape_info_label"].set("Letter made\nCancel manually")
    widgets["root"].update_idletasks()
    os.startfile(letter_path)


def recall():
    widgets["scrape_info_label"].set("")
    widgets["root"].update_idletasks()

    # Navigate to patient info page in BC
    pya.moveTo(100, 450, duration=0.1)
    pya.click()
    pya.press("up", presses=9)
    pya.press("enter")
    time.sleep(0.5)

    # Try scraping up to 3 times
    for attempt in range(3):
        pya.moveTo(EMAIL_POS)
        pya.doubleClick()
        state["email"] = scraper()

        pya.moveTo(MRN_POS)
        pya.doubleClick()
        state["mrn"] = scraper()

        pya.moveTo(DOB_POS)
        pya.doubleClick()
        state["dob"] = scraper()

        mrn_ok = state["mrn"].isdigit()
        dob_ok = parse_dob()

        if mrn_ok and dob_ok:
            break

        widgets["scrape_info_label"].set(f"Scrape failed (attempt {attempt + 1}/3)")
        widgets["root"].update_idletasks()
        time.sleep(0.5)
    else:
        # All 3 attempts failed — skip this patient
        widgets["scrape_info_label"].set("Scrape failed — skipping patient")
        widgets["root"].update_idletasks()
        no_recall()
        return

    # Scrape succeeded — send text
    widgets["scrape_info_label"].set("Sending text")
    widgets["root"].update_idletasks()
    send_text()

    # Decide: email or letter
    if is_email(state["email"]) and state["email"] not in {"", "na"}:
        recall_compose()
    else:
        widgets["scrape_info_label"].set("No email — making letter")
        widgets["root"].update_idletasks()
        letter_compose()


def no_recall():
    write_csv(attended="yes")
    widgets["scrape_info_label"].set("No recall sent -> finish")
    widgets["root"].update_idletasks()
    state["recall_number"] = "none"
    state["recall_type"] = "none"
    pya.click(100, 400)
    pya.press("up", presses=3)
    pya.press("enter")
    pya.hotkey("alt", "n")
    pya.press("enter")

    widgets["patient_label_var"].set("")
    finish_recall()


# --- Patient flow control ---


def collect_file():
    state["full_path"] = askopenfilename()

    widgets["file_label_var"].set("    Open Blue Chip now.")
    widgets["create_datafile_button"].config(state="normal", style="Normal.TButton")
    widgets["root"].update_idletasks()


def extract():
    text2 = ""

    with open(state["full_path"], "r") as rtf_file:
        rtf_content = rtf_file.read()

    text = rtf_to_text(rtf_content)

    for i, line in enumerate(text.splitlines()):
        if i == 0:
            continue
        elif "|||||" in line or not line:
            continue
        else:
            text2 += line
            text2 += "\n"

    for line in text2.splitlines():
        local_list = []
        for i, field in enumerate(line.split("|")):
            if i == 0:
                local_list.append(field)
            elif i == 1:
                doc_name = field.split()[2]
                if doc_name == "Vickers":
                    doc_name = "Mill"
                local_list.append(doc_name)
            elif i == 2:
                local_list.append(field)
            elif i == 3:
                local_list.append(field)
        if local_list != [""]:
            state["output_list_4"].append(local_list)

    state["output_list_4"].reverse()
    state["num_to_do"] = len(state["output_list_4"])
    widgets["count_label_var"].set(f"{state['num_to_do']} patients to do.")
    filename = Path(state["full_path"]).stem
    widgets["file_label_var"].set(f"Working on {filename}.")
    if not args.nopickle:
        set_pickled_list()
    widgets["open_file_button"].config(state="disabled", style="Disabled.TButton")
    widgets["create_datafile_button"].config(state="disabled", style="Disabled.TButton")
    next_patient()


def next_patient():
    if not args.nopickle:
        get_pickled_list()
    try:
        state["pat"] = state["output_list_4"].pop()

    except IndexError:
        widgets["patient_label_var"].set("Finished!")
        state["num_to_do"] = 0
        widgets["count_label_var"].set(f"{state['num_to_do']} patients to do.")
        state["full_path"] = ""
        widgets["open_file_button"].config(state="normal", style="Normal.TButton")
        widgets["create_datafile_button"].config(state="normal", style="Normal.TButton")
        state["output_list_4"] = []

    name = state["pat"][0]
    state["phone"] = state["pat"][2]
    name_for_label = f"{name}\n{state['phone']}"
    widgets["patient_label_var"].set(name_for_label)

    state["num_to_do"] = len(state["output_list_4"])
    widgets["count_label_var"].set(f"{state['num_to_do']} patients to do.")
    widgets["scrape_info_label"].set("")
    widgets["root"].update_idletasks()
    state["phone"] = state["pat"][2].replace("-", "")

    if state["first_run"]:
        pya.alert("Make sure Blue Chip is open then press OK.")
        state["first_run"] = False

    if state["phone"] and state["phone"][0] == "0":
        open_bc_by_phone()
    else:
        open_bc_by_name()


def close_out():
    if not args.nopickle:
        set_pickled_list()
    full_name = state["pat"][0]
    widgets["scrape_info_label"].set(f"{full_name} finished")
    widgets["root"].update_idletasks()
    pya.moveTo(CLOSE_POS[0], CLOSE_POS[1])
    pya.click()
    pya.hotkey("alt", "n")


def finish_recall():
    close_out()
    if state["output_list_4"]:
        widgets["patient_label_var"].set("")
        next_patient()
    else:
        widgets["patient_label_var"].set("Finished")


def finish_exit():
    close_out()
    sys.exit(1)


# --- Menu commands ---


def open_letters():
    os.startfile(LETTERS_PATH)


def reset_program():
    subprocess.run([sys.executable, str(PICKLER_PATH)])
    sys.exit(1)


# --- GUI ---


def main():
    if not args.nopickle:
        get_pickled_list()

    widgets["root"] = Tk()

    widgets["file_label_var"] = StringVar()
    widgets["count_label_var"] = StringVar()
    widgets["patient_label_var"] = StringVar()
    widgets["scrape_info_label"] = StringVar()

    root = widgets["root"]
    root.geometry("350x450+840+50")
    root.title("First Recalls")
    root.option_add("*(tearOff)", FALSE)

    menubar = Menu(root)
    root.config(menu=menubar)
    menu_extras = Menu(menubar)
    menubar.add_cascade(menu=menu_extras, label="Extras")
    menu_extras.add_command(label="Letters folder", command=open_letters)
    menu_extras.add_command(label="Reset Program", command=reset_program)

    button_style = ttk.Style()
    button_style.configure("Normal.TButton", background="lightgray", foreground="blue")
    button_style.configure(
        "Disabled.TButton", background="lightgray", foreground="darkgray"
    )

    mainframe = ttk.Frame(root, padding="3 3 12 12")
    mainframe.grid(column=0, row=0, sticky=(N, W, E, S))
    mainframe.columnconfigure(0, weight=1)
    mainframe.rowconfigure(0, weight=1)

    topframe = Frame(mainframe)
    topframe.grid(column=0, row=0, sticky=(N, W, E, S))
    topframe.columnconfigure(0, weight=1)
    topframe.rowconfigure(0, weight=1)

    midframe = Frame(mainframe, height=2, bg="black")
    midframe.grid(column=0, row=2, sticky=(N, W, E, S))
    midframe.columnconfigure(0, weight=1)
    midframe.rowconfigure(0, weight=1)

    bottomframe = Frame(mainframe)
    bottomframe.grid(column=0, row=3, sticky=(N, W, E, S))
    bottomframe.columnconfigure(0, weight=1)
    bottomframe.rowconfigure(0, weight=1)

    # --- Top panel: file selection ---
    widgets["open_file_button"] = ttk.Button(
        topframe, text="Open Blue Chip File", command=collect_file
    )
    widgets["open_file_button"].grid(column=0, row=0, sticky=W)

    file_label = ttk.Label(topframe, textvariable=widgets["file_label_var"])
    file_label.grid(column=1, row=0, sticky=E)

    widgets["create_datafile_button"] = ttk.Button(
        topframe, text="Create Datafile", command=extract
    )
    widgets["create_datafile_button"].grid(column=0, row=1, sticky=W)

    count_label = ttk.Label(topframe, textvariable=widgets["count_label_var"])
    count_label.grid(column=1, row=1, sticky=E)

    # --- Bottom panel: patient processing ---
    patient_label = ttk.Label(bottomframe, textvariable=widgets["patient_label_var"])
    patient_label.grid(column=0, row=0, sticky=W)

    open_by_name_button = ttk.Button(
        bottomframe, text="Open by name", command=open_bc_by_name
    )
    open_by_name_button.grid(column=1, row=0, sticky=E)

    recall_button = ttk.Button(bottomframe, text="Recall", command=recall)
    recall_button.grid(column=0, row=1, sticky=W)

    no_recall_button = ttk.Button(bottomframe, text="No Recall", command=no_recall)
    no_recall_button.grid(column=1, row=1, sticky=E)

    scrape_label = ttk.Label(bottomframe, textvariable=widgets["scrape_info_label"])
    scrape_label.grid(column=0, row=2, sticky=W)

    finish_new_button = ttk.Button(
        bottomframe, text="Finish & new", command=finish_recall
    )
    finish_new_button.grid(column=0, row=3, sticky=W)

    finish_exit_button = ttk.Button(
        bottomframe, text="Finish & exit", command=finish_exit
    )
    finish_exit_button.grid(column=1, row=3, sticky=E)

    for child in mainframe.winfo_children():
        child.grid_configure(padx=5, pady=10)

    for child in bottomframe.winfo_children():
        child.grid_configure(padx=5, pady=20)

    # --- Initial state ---
    if state["output_list_4"]:
        widgets["open_file_button"].config(state="disabled", style="Disabled.TButton")
        widgets["create_datafile_button"].config(
            state="disabled", style="Disabled.TButton"
        )
        state["num_to_do"] = len(state["output_list_4"])
        widgets["count_label_var"].set(f"{state['num_to_do']} patients to do.")
        next_patient()
    else:
        widgets["open_file_button"].config(state="normal", style="Normal.TButton")
        widgets["create_datafile_button"].config(state="normal", style="Normal.TButton")
        widgets["count_label_var"].set("")

    widgets["scrape_info_label"].set("")

    root.attributes("-topmost", True)
    root.mainloop()


if __name__ == "__main__":
    main()
